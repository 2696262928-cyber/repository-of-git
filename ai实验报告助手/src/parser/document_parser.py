from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile

import fitz
from docx import Document


@dataclass
class ParsedReport:
    file_name: str
    file_type: str
    full_text: str
    code_blocks: list[str]
    tables: list[str]
    page_count: int = 0
    image_count: int = 0


@dataclass
class ImageContext:
    """图片及其在文档中的上下文位置。"""
    image_bytes: bytes
    page_number: int          # 1-based 页码
    surrounding_text: str     # 图片前后的上下文文本
    nearest_section: str      # 最近的章节名


def parse_uploaded_file(uploaded_file) -> ParsedReport:
    suffix = Path(uploaded_file.name).suffix.lower().lstrip(".")
    content = uploaded_file.getvalue()

    if suffix == "docx":
        text, tables, page_count, image_count = _parse_docx(content)
    elif suffix == "pdf":
        text, page_count, image_count = _parse_pdf(content)
        tables = []
    elif suffix in {"txt", "md"}:
        text, tables, page_count, image_count = content.decode("utf-8", errors="ignore"), [], 1, 0
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return ParsedReport(
        file_name=uploaded_file.name,
        file_type=suffix,
        full_text=text.strip(),
        code_blocks=_extract_simple_code_blocks(text),
        tables=tables,
        page_count=page_count,
        image_count=image_count,
    )


# ── 图片提取 ─────────────────────────────────────────────


def extract_document_images(uploaded_file) -> list[bytes]:
    """从上传的文档中提取所有嵌入图片的 bytes。

    Returns:
        list[bytes] — 每个元素是一张图片的 PNG/JPEG 字节数据。
    """
    suffix = Path(uploaded_file.name).suffix.lower().lstrip(".")
    content = uploaded_file.getvalue()

    if suffix == "pdf":
        return _extract_pdf_images(content)
    if suffix == "docx":
        return _extract_docx_images(content)
    return []


def _extract_pdf_images(content: bytes) -> list[bytes]:
    """从 PDF 中提取所有嵌入图片。"""
    images = []
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            for page in document:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        base_image = document.extract_image(xref)
                        images.append(base_image["image"])
                    except Exception:
                        continue
    except Exception:
        pass
    return images


def _extract_docx_images(content: bytes) -> list[bytes]:
    """从 DOCX 中提取所有 word/media/ 下的图片。"""
    images = []
    try:
        with ZipFile(BytesIO(content)) as archive:
            for name in archive.namelist():
                if name.startswith("word/media/") and name.lower().endswith(
                    (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif")
                ):
                    images.append(archive.read(name))
    except Exception:
        pass
    return images


# ── 带上下文的图片提取 ──────────────────────────────────────

# 章节关键词（用于判断图片最近的章节）
_SECTION_KEYWORDS = [
    ("实验目的", ["实验目的", "objective"]),
    ("实验环境", ["实验环境", "environment", "configuration"]),
    ("实验原理", ["实验原理", "principle", "overview", "background"]),
    ("实验步骤", ["实验步骤", "实验过程", "procedure", "step"]),
    ("核心代码", ["核心代码", "source code", "main.c"]),
    ("运行结果", ["运行结果", "实验结果", "results", "result"]),
    ("结果分析", ["结果分析", "analysis", "waveform", "现象分析"]),
    ("实验总结", ["实验总结", "conclusion", "summary"]),
]


def extract_document_images_with_context(uploaded_file) -> list[ImageContext]:
    """提取文档中所有嵌入图片及其上下文位置。

    Returns:
        list[ImageContext] — 每张图片的字节数据 + 页码 + 前后文 + 最近章节。
    """
    suffix = Path(uploaded_file.name).suffix.lower().lstrip(".")
    content = uploaded_file.getvalue()

    if suffix == "pdf":
        return _extract_pdf_images_with_context(content)
    if suffix == "docx":
        return _extract_docx_images_with_context(content)
    return []


def _extract_pdf_images_with_context(content: bytes) -> list[ImageContext]:
    """从 PDF 中提取图片及其上下文（页码 + 周围文本）。"""
    results = []
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            for page_index, page in enumerate(document):
                page_text = page.get_text("text")
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        base_image = document.extract_image(xref)
                        img_bytes = base_image["image"]
                    except Exception:
                        continue

                    # 图片在页面上的位置
                    bbox = page.get_image_bbox(img_info)
                    y_mid = (bbox.y0 + bbox.y1) / 2 if bbox else page.rect.height / 2

                    # 取该页文本中靠近图片位置的前后文
                    context_text = _get_text_near_position(page_text, y_mid, page.rect.height)
                    nearest_section = _find_nearest_section(page_text)

                    results.append(ImageContext(
                        image_bytes=img_bytes,
                        page_number=page_index + 1,
                        surrounding_text=context_text,
                        nearest_section=nearest_section,
                    ))
    except Exception:
        pass
    return results


def _extract_docx_images_with_context(content: bytes) -> list[ImageContext]:
    """从 DOCX 中提取图片及其上下文（段落位置）。"""
    results = []
    try:
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        document = Document(tmp_path)
        # 收集全文段落
        all_paragraphs = [p.text for p in document.paragraphs]
        full_text = "\n".join(all_paragraphs)

        # DOCX 图片通过 inline_shapes 无法直接获取字节，通过 zip 提取
        with ZipFile(BytesIO(content)) as archive:
            image_names = [
                n for n in archive.namelist()
                if n.startswith("word/media/") and n.lower().endswith(
                    (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif")
                )
            ]
            for i, name in enumerate(image_names):
                img_bytes = archive.read(name)
                # DOCX 中图片大致对应全文前 i/len 比例的位置
                ratio = (i + 1) / len(image_names) if image_names else 0
                context_text = full_text[:500] if i == 0 else full_text[-500:]
                nearest_section = _find_nearest_section(full_text)

                results.append(ImageContext(
                    image_bytes=img_bytes,
                    page_number=0,  # DOCX 无页码
                    surrounding_text=context_text,
                    nearest_section=nearest_section,
                ))
    except Exception:
        pass
    return results


def _get_text_near_position(page_text: str, y_mid: float, page_height: float) -> str:
    """获取页面中靠近某垂直位置的文本片段。"""
    if not page_text.strip():
        return ""
    # 取前 300 字符作为上下文
    ratio = y_mid / page_height if page_height > 0 else 0.5
    lines = page_text.splitlines()
    start = int(len(lines) * max(0, ratio - 0.15))
    end = int(len(lines) * min(1, ratio + 0.15))
    context_lines = lines[start:end] if end > start else lines[:5]
    return "\n".join(context_lines)[:300]


def _find_nearest_section(text: str) -> str:
    """在文本中查找最近匹配到的章节名。"""
    text_lower = text.lower()
    for section_name, keywords in reversed(_SECTION_KEYWORDS):
        for kw in keywords:
            if kw in text_lower:
                return section_name
    return "未知章节"


# ── 内部解析 ─────────────────────────────────────────────


def _parse_docx(content: bytes) -> tuple[str, list[str], int, int]:
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    document = Document(tmp_path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            rows.append(" | ".join(cells))
        tables.append("\n".join(rows))

    text_parts = paragraphs + tables
    return "\n\n".join(text_parts), tables, 0, len(document.inline_shapes)


def _parse_pdf(content: bytes) -> tuple[str, int, int]:
    with fitz.open(stream=content, filetype="pdf") as document:
        pages = [page.get_text("text") for page in document]
        image_count = sum(len(page.get_images(full=True)) for page in document)
        page_count = document.page_count
    return "\n\n".join(pages), page_count, image_count


def _extract_simple_code_blocks(text: str) -> list[str]:
    blocks = []
    current = []
    in_block = False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_block and current:
                blocks.append("\n".join(current))
                current = []
            in_block = not in_block
            continue
        if in_block:
            current.append(line)
    return blocks


def render_pdf_page_previews(content: bytes, max_pages: int = 3, zoom: float = 1.2) -> list[bytes]:
    """Render the first PDF pages as PNG bytes for visual preview in the UI."""
    previews = []
    matrix = fitz.Matrix(zoom, zoom)
    with fitz.open(stream=content, filetype="pdf") as document:
        for page in document[:max_pages]:
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            previews.append(pixmap.tobytes("png"))
    return previews
