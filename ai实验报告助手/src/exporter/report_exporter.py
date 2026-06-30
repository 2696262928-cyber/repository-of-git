from datetime import datetime
from io import BytesIO


TASK_LABELS = {
    "report_review": "全面评阅",
    "focused_review": "按需求重点检查",
    "report_qa": "基于报告问答",
    "revision_plan": "生成修改计划",
    "teacher_comment": "生成教师评语",
}


def export_markdown_review(
    file_name: str,
    course_result: dict,
    review: dict,
    knowledge_snippets: list[dict] | None = None,
    task_type: str = "report_review",
    user_instruction: str = "",
) -> str:
    lines = [
        "# 实验报告检测结果",
        "",
        "## 基本信息",
        "",
        f"- 文件名：{file_name}",
        f"- 课程类型：{course_result['course_name']}",
        f"- 任务模式：{_task_label(task_type)}",
        f"- 用户需求：{user_instruction or '未填写'}",
        f"- 检测时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
    ]

    if _has_review_score(review):
        if review.get("focused_conclusion"):
            lines.extend(["## 针对用户需求的结论", "", str(review.get("focused_conclusion", "")), ""])
        _append_scored_review_markdown(lines, review)
    else:
        _append_task_result_markdown(lines, review, task_type)

    _append_knowledge_snippets_markdown(lines, knowledge_snippets or [])
    return "\n".join(lines)


def _append_scored_review_markdown(lines: list[str], review: dict) -> None:
    lines.extend(
        [
            "## 总体评分",
            "",
            f"- 总分：{review.get('total_score', 0)} / 100",
            f"- 总评：{review.get('summary', '')}",
            "",
            "## 分项评分",
            "",
            "| 维度 | 得分 | 满分 | 说明 |",
            "| --- | ---: | ---: | --- |",
        ]
    )
    for item in review.get("dimension_scores", []):
        lines.append(
            f"| {item.get('dimension', '')} | {item.get('score', '')} | {item.get('max_score', '')} | {item.get('comment', '')} |"
        )

    lines.extend(["", "## 主要问题与修改建议", ""])
    problems = review.get("problems", [])
    if not problems:
        lines.append("- 未发现明显问题。")
    for index, problem in enumerate(problems, start=1):
        lines.extend(
            [
                f"### {index}. {problem.get('type', '问题')}",
                "",
                f"- 位置：{problem.get('location', '')}",
                f"- 问题：{problem.get('description', '')}",
                f"- 建议：{problem.get('suggestion', '')}",
                "",
            ]
        )

    lines.extend(["## 教师评语草稿", "", str(review.get("teacher_comment", "")), "", "## 风险提示", ""])
    risk_warnings = review.get("risk_warnings", [])
    if risk_warnings:
        for warning in risk_warnings:
            lines.append(f"- {warning}")
    else:
        lines.append("- 未发现明显风险提示。")


def _append_task_result_markdown(lines: list[str], result: dict, task_type: str) -> None:
    lines.extend(["## 任务结果", ""])
    if task_type == "report_qa":
        lines.extend(["### 回答", "", str(result.get("answer", "")), ""])
        _append_list_markdown(lines, "依据", result.get("evidence", []))
        _append_list_markdown(lines, "缺失信息", result.get("missing_info", []))
        _append_list_markdown(lines, "建议", result.get("suggestions", []))
        _append_list_markdown(lines, "引用来源", result.get("cited_sources", []))
    elif task_type == "revision_plan":
        lines.extend(["### 修改计划总览", "", str(result.get("summary", "")), ""])
        _append_list_markdown(lines, "优先级修改", _format_priority_actions(result.get("priority_actions", [])))
        _append_list_markdown(lines, "快速修复", result.get("quick_fixes", []))
        _append_list_markdown(lines, "深度改进", result.get("deep_improvements", []))
        lines.extend(["### 预计改进收益", "", str(result.get("expected_score_gain", "")), ""])
    elif task_type == "teacher_comment":
        lines.extend(["### 教师评语草稿", "", str(result.get("teacher_comment", "")), ""])
        _append_list_markdown(lines, "主要优点", result.get("strengths", []))
        _append_list_markdown(lines, "主要不足", result.get("weaknesses", []))
        lines.extend(["### 建议分数或等级区间", "", str(result.get("score_suggestion", "")), ""])
        _append_list_markdown(lines, "改进提醒", result.get("improvement_notes", []))
    else:
        lines.extend(["```json", str(result), "```", ""])


def _append_list_markdown(lines: list[str], title: str, items) -> None:
    lines.extend([f"### {title}", ""])
    if isinstance(items, str):
        items = [items] if items.strip() else []
    if not items:
        lines.append("- 暂无内容。")
    else:
        for item in items:
            lines.append(f"- {item}")
    lines.append("")


def _format_priority_actions(actions) -> list[str]:
    formatted = []
    for action in actions or []:
        if isinstance(action, dict):
            formatted.append(
                f"[{action.get('priority', '未标注')}] {action.get('target', '')}：{action.get('action', '')}；原因：{action.get('reason', '')}"
            )
        else:
            formatted.append(str(action))
    return formatted


def export_docx_review(
    file_name: str,
    course_result: dict,
    review: dict,
    knowledge_snippets: list[dict] | None = None,
    task_type: str = "report_review",
    user_instruction: str = "",
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    _setup_document_style(document)

    title = document.add_heading("实验报告智能检测结果", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("基本信息", level=2)
    basic = document.add_table(rows=0, cols=2)
    _style_table(basic)
    _add_kv_row(basic, "文件名", file_name)
    _add_kv_row(basic, "课程类型", course_result.get("course_name", ""))
    _add_kv_row(basic, "任务模式", _task_label(task_type))
    _add_kv_row(basic, "用户需求", user_instruction or "未填写")
    _add_kv_row(basic, "检测时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    if _has_review_score(review):
        if review.get("focused_conclusion"):
            document.add_heading("针对用户需求的结论", level=2)
            document.add_paragraph(str(review.get("focused_conclusion", "")))
        _append_scored_review_docx(document, review)
    else:
        _append_task_result_docx(document, review, task_type)

    _append_knowledge_snippets_docx(document, knowledge_snippets or [])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _append_scored_review_docx(document, review: dict) -> None:
    document.add_heading("总体评分", level=2)
    document.add_paragraph(f"总分：{review.get('total_score', 0)} / 100", style="Intense Quote")
    document.add_paragraph(f"总评：{review.get('summary', '')}")

    document.add_heading("分项评分", level=2)
    score_table = document.add_table(rows=1, cols=4)
    _style_table(score_table)
    headers = ["维度", "得分", "满分", "说明"]
    for cell, header in zip(score_table.rows[0].cells, headers):
        cell.text = header
    _style_header_row(score_table.rows[0])
    for item in review.get("dimension_scores", []):
        row = score_table.add_row().cells
        row[0].text = str(item.get("dimension", ""))
        row[1].text = str(item.get("score", ""))
        row[2].text = str(item.get("max_score", ""))
        row[3].text = str(item.get("comment", ""))

    document.add_heading("主要问题与修改建议", level=2)
    problems = review.get("problems", [])
    if not problems:
        document.add_paragraph("未发现明显问题。")
    for index, problem in enumerate(problems, start=1):
        document.add_heading(f"{index}. {problem.get('type', '问题')}", level=3)
        document.add_paragraph(f"位置：{problem.get('location', '')}")
        document.add_paragraph(f"问题：{problem.get('description', '')}")
        document.add_paragraph(f"建议：{problem.get('suggestion', '')}")

    document.add_heading("教师评语草稿", level=2)
    document.add_paragraph(str(review.get("teacher_comment", "")))

    document.add_heading("风险提示", level=2)
    risk_warnings = review.get("risk_warnings", [])
    if risk_warnings:
        for warning in risk_warnings:
            document.add_paragraph(str(warning), style="List Bullet")
    else:
        document.add_paragraph("未发现明显风险提示。")


def _append_task_result_docx(document, result: dict, task_type: str) -> None:
    document.add_heading("任务结果", level=2)
    if task_type == "report_qa":
        document.add_heading("回答", level=3)
        document.add_paragraph(str(result.get("answer", "")))
        _append_list_docx(document, "依据", result.get("evidence", []))
        _append_list_docx(document, "缺失信息", result.get("missing_info", []))
        _append_list_docx(document, "建议", result.get("suggestions", []))
        _append_list_docx(document, "引用来源", result.get("cited_sources", []))
    elif task_type == "revision_plan":
        document.add_heading("修改计划总览", level=3)
        document.add_paragraph(str(result.get("summary", "")))
        _append_list_docx(document, "优先级修改", _format_priority_actions(result.get("priority_actions", [])))
        _append_list_docx(document, "快速修复", result.get("quick_fixes", []))
        _append_list_docx(document, "深度改进", result.get("deep_improvements", []))
        document.add_heading("预计改进收益", level=3)
        document.add_paragraph(str(result.get("expected_score_gain", "")))
    elif task_type == "teacher_comment":
        document.add_heading("教师评语草稿", level=3)
        document.add_paragraph(str(result.get("teacher_comment", "")))
        _append_list_docx(document, "主要优点", result.get("strengths", []))
        _append_list_docx(document, "主要不足", result.get("weaknesses", []))
        document.add_heading("建议分数或等级区间", level=3)
        document.add_paragraph(str(result.get("score_suggestion", "")))
        _append_list_docx(document, "改进提醒", result.get("improvement_notes", []))
    else:
        document.add_paragraph(str(result))


def _append_list_docx(document, title: str, items) -> None:
    document.add_heading(title, level=3)
    if isinstance(items, str):
        items = [items] if items.strip() else []
    if not items:
        document.add_paragraph("暂无内容。")
        return
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def _has_review_score(result: dict) -> bool:
    return "total_score" in result and "dimension_scores" in result and "problems" in result


def _task_label(task_type: str) -> str:
    return TASK_LABELS.get(task_type, task_type)


def _setup_document_style(document) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = document.styles["Normal"]
    normal.font.name = "等线"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "等线"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.color.rgb = RGBColor(31, 41, 55)


def _style_table(table) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def _style_header_row(row) -> None:
    for cell in row.cells:
        _set_cell_shading(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _append_knowledge_snippets_markdown(lines: list[str], snippets: list[dict]) -> None:
    if not snippets:
        return
    lines.extend(["", "## 本地课程知识库引用", ""])
    for index, item in enumerate(snippets, start=1):
        keywords = "、".join(item.get("matched_keywords", [])[:8]) or "无"
        references = item.get("references", [])
        lines.extend(
            [
                f"### {index}. {item.get('title', '知识片段')}",
                "",
                f"- 来源：{item.get('source', item.get('file', ''))}",
                f"- BM25 分数：{item.get('score', '')}",
                f"- 匹配关键词：{keywords}",
                "",
            ]
        )
        if references:
            lines.append("- 公开参考：")
            lines.extend(f"  - {reference}" for reference in references[:3])
            lines.append("")
        lines.extend([str(item.get("content", "")), ""])


def _append_knowledge_snippets_docx(document, snippets: list[dict]) -> None:
    if not snippets:
        return
    document.add_heading("本地课程知识库引用", level=2)
    for index, item in enumerate(snippets, start=1):
        keywords = "、".join(item.get("matched_keywords", [])[:8]) or "无"
        references = item.get("references", [])
        document.add_heading(f"{index}. {item.get('title', '知识片段')}", level=3)
        document.add_paragraph(f"来源：{item.get('source', item.get('file', ''))}")
        document.add_paragraph(f"BM25 分数：{item.get('score', '')}")
        document.add_paragraph(f"匹配关键词：{keywords}")
        if references:
            document.add_paragraph("公开参考：")
            for reference in references[:3]:
                document.add_paragraph(str(reference), style="List Bullet")
        document.add_paragraph(str(item.get("content", "")))


def _add_kv_row(table, key: str, value: str) -> None:
    row = table.add_row().cells
    row[0].text = key
    row[1].text = str(value)
